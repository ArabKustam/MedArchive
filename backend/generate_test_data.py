# generate_test_data.py
import os
import json
import zipfile
import tempfile
from pathlib import Path
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Change working directory to backend root directory
os.chdir(Path(__file__).parent.absolute())

def create_catalog_json(path: Path):
    services = [
        {
            "service_name": "МРТ головного мозга",
            "category": "МРТ",
            "synonyms": ["Магнитно-резонансная томография головного мозга", "МРТ мозга", "МРТ головы"]
        },
        {
            "service_name": "УЗИ органов брюшной полости",
            "category": "УЗИ",
            "synonyms": ["УЗИ брюшной полости", "УЗИ ОБП", "Ультразвуковое исследование ОБП"]
        },
        {
            "service_name": "Общий анализ крови",
            "category": "Лаборатория",
            "synonyms": ["ОАК", "Клинический анализ крови"]
        },
        {
            "service_name": "Прием врача-терапевта",
            "category": "Консультации",
            "synonyms": ["Консультация терапевта", "Прием терапевта", "Осмотр терапевта"]
        },
        {
            "service_name": "Электрокардиография",
            "category": "Функциональная диагностика",
            "synonyms": ["ЭКГ", "Снятие ЭКГ", "Электрокардиограмма"]
        },
        {
            "service_name": "Рентгенография органов грудной клетки",
            "category": "Рентген",
            "synonyms": ["Рентген легких", "Флюорография", "Рентген грудной клетки"]
        },
        {
            "service_name": "Прием врача-кардиолога",
            "category": "Консультации",
            "synonyms": ["Консультация кардиолога", "Прием кардиолога"]
        },
        {
            "service_name": "УЗИ сердца",
            "category": "УЗИ",
            "synonyms": ["Эхокардиография", "ЭхоКГ"]
        },
        {
            "service_name": "Анализ крови на глюкозу",
            "category": "Лаборатория",
            "synonyms": ["Глюкоза крови", "Сахар крови"]
        },
        {
            "service_name": "КТ легких",
            "category": "КТ",
            "synonyms": ["Компьютерная томография легких", "КТ органов грудной клетки"]
        }
    ]
    path.write_text(json.dumps(services, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Created catalog file at: {path}")

def generate_xlsx(path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Прайс-лист"
    
    # Headers
    headers = ["Код услуги", "Наименование услуги", "Цена для резидентов (KZT)", "Цена для нерезидентов (KZT)"]
    ws.append(headers)
    
    # Rows
    rows = [
        ["001", "МРТ головного мозга", "18000", "25000"],
        ["002", "УЗИ органов брюшной полости", "7000", "10000"],
        ["003", "Общий анализ крови", "3000", "4500"],
        ["004", "ЭКГ с расшифровкой", "4000", "6000"],  # Will fuzzy match "Электрокардиография"
        ["005", "Услуга ХХХ (Неизвестная)", "12000", "15000"]  # Will need manual review
    ]
    for r in rows:
        ws.append(r)
    wb.save(path)
    print(f"Generated Excel price list at: {path}")

def generate_docx(path: Path):
    doc = Document()
    doc.add_heading("Тарифы клиники Партнер Б", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Наименование услуги"
    hdr_cells[1].text = "Стоимость (резиденты)"
    hdr_cells[2].text = "Стоимость (нерезиденты)"
    
    data = [
        ("Консультация терапевта", "6000", "8500"),  # Will match synonym of "Прием врача-терапевта"
        ("Рентген легких", "5000", "7000"),  # Will match synonym of "Рентгенография органов грудной клетки"
        ("УЗИ почек (Неизвестная)", "8000", "11000")  # Will need manual review
    ]
    
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        
    doc.save(path)
    print(f"Generated DOCX price list at: {path}")

def generate_pdf(path: Path):
    # Register Arial for Cyrillic support
    windir = os.environ.get("WINDIR", "C:\\Windows")
    arial_path = os.path.join(windir, "Fonts", "arial.ttf")
    if os.path.exists(arial_path):
        pdfmetrics.registerFont(TTFont("Arial", arial_path))
        font_name = "Arial"
    else:
        font_name = "Helvetica"
        
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    story = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=18,
        leading=22,
        spaceAfter=15
    )
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=12
    )
    
    story.append(Paragraph("Прайс-лист клиники Партнер С", title_style))
    
    data = [
        [Paragraph("Наименование", cell_style), Paragraph("Резидент", cell_style), Paragraph("Нерезидент", cell_style)],
        [Paragraph("Эхокардиография", cell_style), Paragraph("15000", cell_style), Paragraph("20000", cell_style)],  # Matches synonym of "УЗИ сердца"
        [Paragraph("Сахар крови", cell_style), Paragraph("2000", cell_style), Paragraph("3000", cell_style)],  # Matches synonym of "Анализ крови на глюкозу"
        [Paragraph("КТ легких", cell_style), Paragraph("22000", cell_style), Paragraph("30000", cell_style)]  # Matches "КТ легких"
    ]
    
    t = Table(data, colWidths=[250, 100, 100])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), '#eeeeee'),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, '#cccccc')
    ]))
    story.append(t)
    
    doc.build(story)
    print(f"Generated PDF price list at: {path}")

def main():
    # Make sure backend/data exists
    data_dir = Path("data")
    data_dir.mkdir(exist_ok=True)
    
    catalog_path = data_dir / "catalog.json"
    create_catalog_json(catalog_path)
    
    # Import seed function
    from app.seed import seed_from_json
    from app.db import init_db, SessionLocal
    from app.models import Service
    
    init_db()
    db = SessionLocal()
    
    try:
        count = db.query(Service).count()
    except Exception:
        count = 0

    # Check if catalog is already seeded to prevent duplicate catalog entries
    if count == 0:
        print("Database catalog is empty. Seeding...")
        n = seed_from_json(catalog_path)
        print(f"Successfully seeded {n} services to database.")
    else:
        print(f"Database already contains {count} services. Skipping seed.")
    db.close()
    
    # Generate test files in temporary directory, and pack into zip
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        
        xlsx_file = tmp_path / "Прайс_Партнер_А_2026.xlsx"
        docx_file = tmp_path / "Тарифы_Партнер_Б_2026.docx"
        pdf_file = tmp_path / "Прайс_Партнер_С_2026.pdf"
        
        generate_xlsx(xlsx_file)
        generate_docx(docx_file)
        generate_pdf(pdf_file)
        
        # Package into zip in root workspace
        zip_out = Path("..") / "test_prices.zip"
        with zipfile.ZipFile(zip_out, "w") as zf:
            zf.write(xlsx_file, xlsx_file.name)
            zf.write(docx_file, docx_file.name)
            zf.write(pdf_file, pdf_file.name)
            
        print(f"\nCreated test price ZIP archive successfully at: {zip_out.resolve()}")

if __name__ == "__main__":
    main()
