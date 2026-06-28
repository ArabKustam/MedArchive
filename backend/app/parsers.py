"""Document parsers: PDF (text + OCR fallback), DOCX, XLSX/XLS.

Intelligent header detection, column mapping, service code extraction,
merged-cell filtering, and robust price parsing.
"""
from __future__ import annotations

import os
import re
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

logger = logging.getLogger(__name__)

# Price pattern: matches numbers like "14 400", "3980", "2 210,00", "15800.50"
# Price pattern: matches numbers like "14 400", "3980", "2 210,00", "15800.50", "18 00С", "36 00("
PRICE_RE = re.compile(
    r"\b(\d[\d\s\u00a0]{0,12}(?:[.,]\d{1,2})?[\u043e\u041e\u0441\u0421oOcС\(\)]?)\b(?!\w)"
    r"|\b(\d[\d\s\u00a0]{0,12}(?:[.,]\d{1,2})?)\s*(тг|тенге|kzt|₸)?",
    re.IGNORECASE,
)

# Service code patterns found in real documents
CODE_RE = re.compile(
    r"\b([A-ZА-Я]{1,4}[\d.\-]{1,15}\d)\b"
    r"|\b(\d{2,4}\.\d{1,4}(?:\.\d+)*)\b",
    re.IGNORECASE,
)

# Units of measurement regex to avoid classifying measurement units as service names
UNIT_RE = re.compile(
    r"^(?:\d+[\d\s\u00a0,.-]*)?\s*(?:операци|процедур|исследован|визит|посещен|услуг|сеанс|единиц|единица|ед|зуб|элемент|зона|блок|курс|проба|область|точка|сустав|челюсть|койко|инъекци|анализ)[а-я]*$",
    re.IGNORECASE,
)

MAX_REASONABLE_PRICE = 10_000_000  # 10 million KZT max
MIN_REASONABLE_PRICE = 5  # Below this, likely a row-number artifact


def _is_garbled_text(text: str) -> bool:
    """Detect garbled PDF text from non-standard font encodings.

    Many Kazakh medical-clinic PDFs embed CIDFont / Type1 subsets that
    pdfplumber cannot decode.  Symptoms include box-drawing characters
    (U+2500–U+25FF), REPLACEMENT CHARACTER (U+FFFD), and a suspiciously
    low ratio of Cyrillic letters in what should be Russian text.
    """
    if not text or len(text.strip()) < 50:
        return True

    # Box-drawing / block-element / replacement characters → almost certainly garbled
    garbage = sum(
        1 for c in text
        if c == '\ufffd' or '\u2500' <= c <= '\u25ff'
    )
    if garbage >= 2:
        return True

    alpha_chars = sum(1 for c in text if c.isalpha())
    if alpha_chars < 20:
        return True

    cyrillic = sum(1 for c in text if '\u0400' <= c <= '\u04ff')
    # A proper Russian medical PDF should be > 40 % Cyrillic among letters
    if cyrillic / alpha_chars < 0.35:
        return True

    return False


def _setup_google_credentials():
    if "GOOGLE_APPLICATION_CREDENTIALS" not in os.environ:
        for parent_level in (2, 1):
            key = Path(__file__).resolve().parents[parent_level] / "consummate-sled-493518-m6-4c991329b233.json"
            if key.exists():
                os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(key)
                break


@dataclass
class ExtractedRow:
    service_name_raw: str
    price_resident_kzt: Optional[float] = None
    price_nonresident_kzt: Optional[float] = None
    service_code_source: Optional[str] = None
    currency_original: str = "KZT"


# ──────────────── Header / column detection ────────────────

_NAME_KW_STEMS = {"наимен", "назван", "услуг", "процедур", "описан", "обследов", "работ"}
_PRICE_KW_STEMS = {"стоимост", "цен", "тариф", "тенг", "kzt", "руб"}

_NAME_KW = {"наименование", "название", "услуга", "услуги", "вид услуги", "медицинские услуги", "процедура", "описание", "обследований", "работа", "работ"}
_PRICE_KW = {"стоимость", "цена", "тариф", "тенге", "kzt", "руб"}
_CODE_KW = {"код", "code", "артикул", "шифр", "мкб", "тарификатор"}
_RESIDENT_KW = {"резидент", "резидентов", "граждан рк", "граждане рк", "республики казахстан", "кандас", "физ. лица", "физические лица", "физлица", "основная цена", "цена рк", "стоимость рк", "основная"}
_NONRESIDENT_KW = {"нерезидент", "нерезидентов", "иностран", "иностранные", "иностранцев", "снг", "ближнего зарубежья", "дальнего зарубежья", "страхов", "дмс", "цена нерезидент", "стоимость нерезидент"}
_SKIP_KW = {
    "наименование", "название", "услуга", "услуги", "стоимость", "цена",
    "резидент", "нерезидент", "код", "номер", "№", "категория", "тариф",
    "п/п", "пп", "ед.изм", "ед. изм", "единица", "кол-во", "количество",
    "итого", "всего", "subtotal", "total", "раздел",
    "приложение", "приказ", "утверждаю", "согласовано", "председатель",
    "правления", "заместитель", "прейскурант",
}


@dataclass
class ColumnMap:
    """Indices of key columns discovered from header rows."""
    name_idx: int = -1
    code_idx: int = -1
    price_res_idx: int = -1      # resident / main price
    price_nonres_idx: int = -1   # non-resident price (may be -1)
    header_row_idx: int = -1     # start row index of header
    header_end_row_idx: int = -1 # end row index of header


def _cell_lower(s: str) -> str:
    return s.strip().lower()


def _is_header_cell_name(s: str) -> bool:
    v = _cell_lower(s)
    if "код" in v or "прейскурант" in v:
        return False
    return any(kw in v for kw in _NAME_KW_STEMS)


def _is_header_cell_price(s: str) -> bool:
    v = _cell_lower(s)
    if "тарификатор" in v:
        return False
    return any(kw in v for kw in _PRICE_KW_STEMS)


def _is_header_cell_code(s: str) -> bool:
    v = _cell_lower(s)
    return any(kw in v for kw in _CODE_KW) or "код" in v or "мкб" in v


def _is_skip_row(cells: List[str]) -> bool:
    """Return True if this row is a header, section title, or merged decoration."""
    non_empty = [c for c in cells if c.strip()]
    if not non_empty:
        return True
    if len(set(c.strip() for c in non_empty)) == 1 and len(non_empty) > 1:
        return True
    combined = " ".join(non_empty).lower()
    if any(kw in combined for kw in ("приложение", "приказ", "утверждаю", "согласовано",
                                      "председатель", "правления", "прейскурант", "оказываемые", "страховых компаний")):
        return True
    if len(non_empty) == 1:
        txt = non_empty[0].strip()
        if txt.isupper() or txt.lower().startswith("раздел"):
            return True
    return False


def _detect_columns(rows: List[List[str]], max_scan: int = 40) -> ColumnMap:
    """Scan rows to find multi-level headers and map key columns."""
    cm = ColumnMap()
    n_rows = min(max_scan, len(rows))
    for ri in range(n_rows):
        cells = [str(c).strip() if c is not None else "" for c in rows[ri]]
        non_empty = [c for c in cells if c.strip()]
        if len(non_empty) < 2:
            continue

        for header_height in range(1, 7):
            if ri + header_height > len(rows):
                continue
            
            num_cols = max(len(r) for r in rows[ri:ri + header_height])
            combined_cols = []
            for ci in range(num_cols):
                col_parts = []
                for h_row in rows[ri:ri + header_height]:
                    if ci < len(h_row) and h_row[ci] is not None:
                        val = str(h_row[ci]).strip()
                        if val and val not in col_parts:
                            col_parts.append(val)
                combined_cols.append(" ".join(col_parts))

            if len([c for c in combined_cols if c.strip()]) < 2:
                continue

            has_name = any(_is_header_cell_name(c) for c in combined_cols)
            has_price = any(_is_header_cell_price(c) or any(kw in c.lower() for kw in _RESIDENT_KW | _NONRESIDENT_KW) for c in combined_cols)

            if has_name and has_price:
                cm = ColumnMap()
                cm.header_row_idx = ri
                cm.header_end_row_idx = ri + header_height - 1

                for ci, c in enumerate(combined_cols):
                    low = c.lower()
                    if cm.code_idx == -1 and _is_header_cell_code(c):
                        cm.code_idx = ci
                    elif cm.name_idx == -1 and _is_header_cell_name(c):
                        cm.name_idx = ci
                    else:
                        is_nonres = any(kw in low for kw in _NONRESIDENT_KW)
                        is_res = any(kw in low for kw in _RESIDENT_KW)
                        is_price = _is_header_cell_price(c)
                        is_without_vat = "без учета ндс" in low or "без ндс" in low

                        if is_nonres and cm.price_nonres_idx == -1:
                            cm.price_nonres_idx = ci
                        elif is_res and (cm.price_res_idx == -1 or "без учета ндс" in combined_cols[cm.price_res_idx].lower()):
                            cm.price_res_idx = ci
                        elif is_price and not is_without_vat:
                            if cm.price_res_idx == -1:
                                cm.price_res_idx = ci
                            elif cm.price_nonres_idx == -1 and ci != cm.price_res_idx:
                                cm.price_nonres_idx = ci

                if cm.name_idx >= 0 and cm.price_res_idx == -1:
                    data_start = cm.header_end_row_idx + 1
                    for dr in rows[data_start: data_start + 5]:
                        dcells = [str(c).strip() if c is not None else "" for c in dr]
                        for ci, c in enumerate(dcells):
                            if ci < len(combined_cols):
                                c_header = combined_cols[ci].lower()
                                if any(skip in c_header for skip in ("№", "п/п", "ед", "изм", "кол")):
                                    continue
                            if ci != cm.name_idx and ci != cm.code_idx and _to_number(c) is not None:
                                if cm.price_res_idx == -1:
                                    cm.price_res_idx = ci
                                elif cm.price_nonres_idx == -1 and ci != cm.price_res_idx:
                                    cm.price_nonres_idx = ci
                        if cm.price_res_idx >= 0:
                            break
                if cm.name_idx >= 0 and cm.price_res_idx >= 0:
                    return cm
    return cm


def _merge_multiline_table_rows(raw_rows: List[List[str]]) -> List[List[str]]:
    """Merge wrapped rows in tabular data where service names or details span multiple physical rows."""
    if not raw_rows:
        return []
    merged = []
    curr_row = None

    for row in raw_rows:
        if not any(c for c in row if c and str(c).strip()):
            continue
        col0 = str(row[0]).strip() if len(row) > 0 else ""
        has_index = bool(re.match(r"^\d{1,5}$", col0))
        has_code = any(_is_code(str(c)) for c in row if c)

        if has_index or has_code or curr_row is None:
            if curr_row:
                merged.append(curr_row)
            curr_row = [str(c or "").strip() for c in row]
        else:
            for ci, val in enumerate(row):
                val_str = str(val or "").strip()
                if not val_str:
                    continue
                if ci < len(curr_row):
                    if curr_row[ci]:
                        # Don't concatenate two numeric values — that would
                        # produce garbage prices like "1340" + "320" → 1340320
                        if _to_number(val_str) is not None and _to_number(curr_row[ci]) is not None:
                            continue
                        if val_str not in curr_row[ci]:
                            curr_row[ci] = curr_row[ci] + " " + val_str
                    else:
                        curr_row[ci] = val_str
                else:
                    while len(curr_row) < ci:
                        curr_row.append("")
                    curr_row.append(val_str)
    if curr_row:
        merged.append(curr_row)
    return merged


# ──────────────── Number / price helpers ────────────────

def _to_number(s: str) -> Optional[float]:
    """Parse a string as a number, handling spaces, commas, and OCR artifacts."""
    if not s:
        return None
    orig = str(s).strip()
    if not orig:
        return None
    # Cells may contain multiple newline-separated values (e.g. "780\n780")
    # from pdfplumber's text-strategy table merging — only parse the first.
    if "\n" in orig:
        orig = orig.split("\n")[0].strip()
        if not orig:
            return None
    cleaned = orig.replace("\u00a0", "").replace(" ", "").replace(",", ".").rstrip(".")
    try:
        val = float(cleaned)
        if 0 <= val <= MAX_REASONABLE_PRICE:
            return val
    except ValueError:
        pass

    t = orig.replace("\u00a0", " ").strip()
    t = re.sub(r"^[\u0437\u0417](?=[\d\u043e\u041e\u0441\u0421oOcС\s])", "3", t)
    t = re.sub(r"(?<=[\d\u0437\u0417])[\u043e\u041e\u0441\u0421oOcС\(\)]", "0", t)
    t = re.sub(r"[\u043e\u041e\u0441\u0421oOcС\(\)](?=[\d\u0437\u0417])", "0", t)
    t = re.sub(r"(?<=\d)[\u043e\u041e\u0441\u0421oOcС\(\)]", "0", t)
    t = re.sub(r"[\u043e\u041e\u0441\u0421oOcС\(\)](?=\d)", "0", t)
    t = re.sub(r"(\d+[\d\s]*)[\u043e\u041e\u0441\u0421oOcС\(\)]+$", r"\g<1>0", t)

    t_clean = re.sub(r"[\s]+", "", t).replace(",", ".").rstrip(".")
    try:
        val = float(t_clean)
        if 0 <= val <= MAX_REASONABLE_PRICE:
            return val
    except ValueError:
        pass
    return None


def _is_code(s: str) -> bool:
    """Check if a string looks like a service code."""
    return bool(CODE_RE.match(s.strip()))


def _extract_row_mapped(cells: List[str], cm: ColumnMap) -> Optional[ExtractedRow]:
    """Extract a row using pre-detected column mapping."""
    if len(cells) <= cm.name_idx:
        return None
    name = cells[cm.name_idx].strip() if cm.name_idx >= 0 else ""
    if not name or len(name) < 2 or UNIT_RE.match(name):
        return None
    if _is_skip_row(cells):
        return None
    if _to_number(name) is not None:
        return None

    code = cells[cm.code_idx].strip() if 0 <= cm.code_idx < len(cells) else None
    if code and not code.strip():
        code = None

    res_price = None
    non_price = None
    if 0 <= cm.price_res_idx < len(cells):
        res_price = _to_number(cells[cm.price_res_idx])
    if 0 <= cm.price_nonres_idx < len(cells):
        non_price = _to_number(cells[cm.price_nonres_idx])

    # Filter row-number artifacts mistaken for prices
    if res_price is not None and res_price < MIN_REASONABLE_PRICE:
        res_price = None
    if non_price is not None and non_price < MIN_REASONABLE_PRICE:
        non_price = None

    if res_price is None and non_price is None:
        return None

    return ExtractedRow(
        service_name_raw=name,
        price_resident_kzt=res_price,
        price_nonresident_kzt=non_price,
        service_code_source=code,
    )


def _extract_row_heuristic(cells: List[str]) -> Optional[ExtractedRow]:
    """Fallback: extract a row without column mapping using heuristics."""
    if _is_skip_row(cells):
        return None

    name = None
    code = None
    nums: List[float] = []

    for ci, c in enumerate(cells):
        text = c.strip()
        if not text:
            continue
        n = _to_number(text)
        if n is not None:
            if ci == 0 and n < 10000:
                continue
            if n >= 10:
                nums.append(n)
            continue
        if _is_code(text):
            code = text
            continue
        if UNIT_RE.match(text):
            continue
        if any(kw in text.lower() for kw in _SKIP_KW):
            continue
        if any(ch.isalpha() for ch in text) and len(text) >= 2:
            if name is None or len(text) > len(name):
                name = text

    if not name or not nums:
        return None

    # Discard obvious row-number artefacts
    nums = [n for n in nums if n >= MIN_REASONABLE_PRICE]
    if not nums:
        return None

    res = nums[-2] if len(nums) >= 2 else nums[0]
    non = nums[-1] if len(nums) >= 2 else None

    return ExtractedRow(
        service_name_raw=name,
        price_resident_kzt=res,
        price_nonresident_kzt=non,
        service_code_source=code,
    )


# ──────────────── PDF (text + tables) ────────────────

def parse_pdf(path: Path) -> List[ExtractedRow]:
    rows: List[ExtractedRow] = []
    logger.info(f"📄 [PDF Engine] Inspecting file: {path.name}")
    try:
        import pdfplumber
    except ImportError:
        logger.error(f"❌ [PDF Engine] pdfplumber is missing! Cannot parse {path.name}")
        return rows

    try:
        with pdfplumber.open(str(path)) as pdf:
            logger.info(f"📄 [PDF Engine] Document opened successfully: {path.name} ({len(pdf.pages)} pages)")
            all_text_pages = []
            last_known_cm: Optional[ColumnMap] = None

            for page in pdf.pages:
                tables = page.extract_tables() or []
                if not tables:
                    tables = page.extract_tables(table_settings={"vertical_strategy": "text", "horizontal_strategy": "text", "snap_tolerance": 3}) or []

                page_table_rows = []
                for table in tables:
                    raw_rows = []
                    for r in table:
                        if not r:
                            continue
                        raw_rows.append([(c or "").strip() for c in r])

                    if not raw_rows:
                        continue

                    merged_rows = _merge_multiline_table_rows(raw_rows)
                    cm = _detect_columns(merged_rows)
                    
                    if cm.name_idx >= 0 and cm.price_res_idx >= 0:
                        last_known_cm = cm
                    elif last_known_cm and last_known_cm.name_idx >= 0 and len(merged_rows[0]) >= last_known_cm.name_idx + 1:
                        cm = last_known_cm

                    start = cm.header_end_row_idx + 1 if cm.header_end_row_idx >= 0 else 0

                    for cells in merged_rows[start:]:
                        if cm.name_idx >= 0:
                            item = _extract_row_mapped(cells, cm)
                        else:
                            item = _extract_row_heuristic(cells)
                        if item:
                            page_table_rows.append(item)

                if page_table_rows:
                    rows.extend(page_table_rows)
                else:
                    t = page.extract_text()
                    if t and len(t.strip()) > 20:
                        all_text_pages.append(t)

            if all_text_pages:
                logger.info(f"📄 [PDF Engine] Extracted {len(rows)} table rows; parsing {len(all_text_pages)} text pages line-by-line...")
                combined_text = "\n".join(all_text_pages)
                rows.extend(_parse_text_lines(combined_text))

            logger.info(f"✅ [PDF Engine Completed] {path.name}: Total extracted positions = {len(rows)}")
    except Exception as e:
        logger.error(f"❌ [PDF Engine Exception] Failed parsing {path.name}: {e}", exc_info=True)
        return rows
    return rows


# ──────────────── OCR (Google Vision + Tesseract) ────────────────

def _ocr_image_google_vision(img_bytes: bytes) -> str:
    _setup_google_credentials()
    try:
        from google.cloud import vision
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=img_bytes)
        response = client.document_text_detection(image=image)
        if not response.full_text_annotation:
            return ""

        words = []
        for page_res in response.full_text_annotation.pages:
            for block in page_res.blocks:
                for paragraph in block.paragraphs:
                    for word in paragraph.words:
                        w_text = "".join([symbol.text for symbol in word.symbols])
                        vertices = word.bounding_box.vertices
                        min_y = min(v.y for v in vertices if v.y is not None)
                        max_y = max(v.y for v in vertices if v.y is not None)
                        min_x = min(v.x for v in vertices if v.x is not None)
                        words.append({"text": w_text, "y": (min_y + max_y) / 2, "x": min_x})

        if not words:
            return response.full_text_annotation.text or ""

        # Group words into table rows by Y coordinate clustering
        words.sort(key=lambda w: w["y"])
        lines = []
        curr_line = []
        curr_y = None
        for w in words:
            if curr_y is None or abs(w["y"] - curr_y) < 14:
                curr_line.append(w)
                curr_y = w["y"] if curr_y is None else sum(item["y"] for item in curr_line) / len(curr_line)
            else:
                curr_line.sort(key=lambda item: item["x"])
                lines.append(" ".join(item["text"] for item in curr_line))
                curr_line = [w]
                curr_y = w["y"]
        if curr_line:
            curr_line.sort(key=lambda item: item["x"])
            lines.append(" ".join(item["text"] for item in curr_line))

        return "\n".join(lines)
    except Exception as e:
        logger.warning(f"Google Cloud Vision OCR warning: {e}")
    return ""


def _ocr_image_to_table_rows(img_bytes: bytes) -> List[List[str]]:
    """Reconstruct table rows from Google Vision word-level bounding boxes.

    Instead of treating OCR output as flat text, this function clusters
    words into *rows* (by Y-centre proximity) and then splits each row
    into *cells* (by large X-gaps between neighbouring words).  The
    resulting 2-D list can be fed straight into ``_detect_columns`` /
    ``_extract_row_mapped``.
    """
    _setup_google_credentials()
    try:
        from google.cloud import vision
    except ImportError:
        return []

    try:
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=img_bytes)
        response = client.document_text_detection(image=image)
        if not response.full_text_annotation:
            return []

        # ── collect words with positions ──
        words: list[dict] = []
        for page_res in response.full_text_annotation.pages:
            for block in page_res.blocks:
                for paragraph in block.paragraphs:
                    for word in paragraph.words:
                        w_text = "".join(sym.text for sym in word.symbols)
                        verts = word.bounding_box.vertices
                        xs = [v.x for v in verts if v.x is not None]
                        ys = [v.y for v in verts if v.y is not None]
                        if not xs or not ys:
                            continue
                        words.append({
                            "text": w_text,
                            "x_min": min(xs),
                            "x_max": max(xs),
                            "y_center": (min(ys) + max(ys)) / 2,
                            "height": max(ys) - min(ys),
                        })

        if len(words) < 5:
            return []

        # ── step 1: cluster words into rows by Y ──
        words.sort(key=lambda w: w["y_center"])
        avg_h = sum(w["height"] for w in words) / len(words)
        row_gap = max(avg_h * 0.55, 6)

        lines: list[list[dict]] = []
        cur_line = [words[0]]
        cur_y = words[0]["y_center"]
        for w in words[1:]:
            if abs(w["y_center"] - cur_y) <= row_gap:
                cur_line.append(w)
                cur_y = sum(x["y_center"] for x in cur_line) / len(cur_line)
            else:
                lines.append(cur_line)
                cur_line = [w]
                cur_y = w["y_center"]
        if cur_line:
            lines.append(cur_line)

        # ── step 2: determine column-gap threshold ──
        all_gaps: list[float] = []
        for line_w in lines:
            line_w.sort(key=lambda w: w["x_min"])
            for i in range(1, len(line_w)):
                g = line_w[i]["x_min"] - line_w[i - 1]["x_max"]
                if g > 2:
                    all_gaps.append(g)
        if not all_gaps:
            return []

        all_gaps.sort()
        p75 = all_gaps[int(len(all_gaps) * 0.75)]
        col_gap = max(p75 * 1.3, 20)

        # ── step 3: split each line into cells ──
        table: list[list[str]] = []
        for line_w in lines:
            line_w.sort(key=lambda w: w["x_min"])
            cells: list[str] = []
            cell_words = [line_w[0]]
            for i in range(1, len(line_w)):
                g = line_w[i]["x_min"] - line_w[i - 1]["x_max"]
                if g > col_gap:
                    cells.append(" ".join(w["text"] for w in cell_words))
                    cell_words = [line_w[i]]
                else:
                    cell_words.append(line_w[i])
            cells.append(" ".join(w["text"] for w in cell_words))
            table.append(cells)

        return table

    except Exception as e:
        logger.warning("Structured OCR table reconstruction failed: %s", e)
        return []


def _parse_text_lines(text: str) -> List[ExtractedRow]:
    """Parse OCR or plain text output line-by-line."""
    rows: List[ExtractedRow] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        low = line.lower()
        if any(kw in low for kw in ("приложение", "приказ", "утверждаю", "согласовано", "председатель", "правления", "прейскурант", "оказываемые")):
            continue

        # Extract service code if present anywhere in line
        code = None
        code_m = CODE_RE.search(line)
        if code_m:
            code = code_m.group()

        # Find all candidate prices in the line
        matches = list(PRICE_RE.finditer(line))
        valid_prices = []
        for m in matches:
            v = _to_number(m.group(1))
            if v is not None and v >= 10:
                val_int = int(v)
                has_curr = bool(m.group(2))
                if 2020 <= val_int <= 2030 and not has_curr:
                    continue
                # Make sure this match is not part of a service code
                if code_m and not (m.start() >= code_m.end() or m.end() <= code_m.start()):
                    continue
                valid_prices.append((m.start(), m.end(), v))

        if not valid_prices:
            continue

        # Service name is extracted between service code (or line start) and the first price
        first_price_idx = valid_prices[0][0]
        if code_m and code_m.start() < first_price_idx:
            raw_name = line[code_m.end():first_price_idx].strip(" .-—:;\t|")
        else:
            raw_name = line[:first_price_idx].strip(" .-—:;\t|")

        # Strip leading row numbers or indexes (e.g. "29 ", "1.1 ")
        cleaned_name = re.sub(r"^(?:[A-ZА-Я]{1,4}[\d.\-]{1,15}\d|\d{2,4}\.\d{1,4}(?:\.\d+)*|\d+[\d.\-]*)\s+", "", raw_name, flags=re.I).strip(" .-—:;\t|")
        name = cleaned_name if len(cleaned_name) >= 2 else raw_name

        if len(name) < 2 or not any(c.isalpha() for c in name):
            continue

        res_price = valid_prices[-2][2] if len(valid_prices) >= 2 else valid_prices[0][2]
        non_price = valid_prices[-1][2] if len(valid_prices) >= 2 else None

        rows.append(ExtractedRow(
            service_name_raw=name,
            price_resident_kzt=res_price,
            price_nonresident_kzt=non_price,
            service_code_source=code,
        ))
    return rows


def parse_pdf_ocr(path: Path) -> List[ExtractedRow]:
    """OCR for scanned / garbled PDFs with structured table reconstruction.

    Renders each page as an image, calls Google Vision API, and rebuilds
    the table grid from word-level bounding boxes before applying the
    standard column-detection + row-extraction pipeline.
    """
    rows: List[ExtractedRow] = []
    images = []

    # Render PDF pages to images
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(str(path))
        for page in pdf:
            img = page.render(scale=2).to_pil()
            images.append(img)
    except Exception:
        images = []

    if not images:
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(str(path), dpi=200)
        except Exception:
            return rows

    import io
    last_known_cm: Optional[ColumnMap] = None

    for img in images:
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        img_bytes = buf.getvalue()
        page_items: List[ExtractedRow] = []

        # Strategy 1: structured table reconstruction from OCR word positions
        table_rows = _ocr_image_to_table_rows(img_bytes)
        if table_rows and len(table_rows) >= 3:
            merged = _merge_multiline_table_rows(table_rows)
            cm = _detect_columns(merged)

            if cm.name_idx >= 0 and cm.price_res_idx >= 0:
                last_known_cm = cm
            elif (
                last_known_cm
                and last_known_cm.name_idx >= 0
                and merged
                and len(merged[0]) >= last_known_cm.name_idx + 1
            ):
                cm = last_known_cm

            start = cm.header_end_row_idx + 1 if cm.header_end_row_idx >= 0 else 0
            for cells in merged[start:]:
                if cm.name_idx >= 0:
                    item = _extract_row_mapped(cells, cm)
                else:
                    item = _extract_row_heuristic(cells)
                if item:
                    page_items.append(item)

        # Strategy 2: flat-text fallback
        if not page_items:
            text = _ocr_image_google_vision(img_bytes)
            if not text:
                try:
                    import pytesseract
                    text = pytesseract.image_to_string(img, lang="rus+eng")
                except Exception:
                    text = ""
            if text:
                page_items = _parse_text_lines(text)

        rows.extend(page_items)

    return rows


def parse_image_file(path: Path) -> List[ExtractedRow]:
    """Extract price items from an image file (JPG/PNG) using Google Vision API."""
    rows: List[ExtractedRow] = []
    try:
        img_bytes = path.read_bytes()
        text = _ocr_image_google_vision(img_bytes)
        if not text:
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(path)
                text = pytesseract.image_to_string(img, lang="rus+eng")
            except Exception:
                text = ""
        rows.extend(_parse_text_lines(text))
    except Exception as e:
        logger.warning(f"Failed to parse image file {path}: {e}")
    return rows


# ──────────────── DOCX ────────────────

def parse_docx(path: Path) -> List[ExtractedRow]:
    rows: List[ExtractedRow] = []
    try:
        from docx import Document
    except ImportError:
        return rows
    try:
        doc = Document(str(path))
    except Exception:
        return rows

    for table in doc.tables:
        all_rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            all_rows.append(cells)

        merged_rows = _merge_multiline_table_rows(all_rows)
        cm = _detect_columns(merged_rows)
        start = cm.header_end_row_idx + 1 if cm.header_end_row_idx >= 0 else 0

        for cells in merged_rows[start:]:
            if cm.name_idx >= 0:
                item = _extract_row_mapped(cells, cm)
            else:
                item = _extract_row_heuristic(cells)
            if item:
                rows.append(item)
    return rows


# ──────────────── XLSX / XLS ────────────────

def _extract_from_sheet_rows(all_rows: List[List[str]]) -> List[ExtractedRow]:
    rows: List[ExtractedRow] = []
    cm = _detect_columns(all_rows)
    start = cm.header_end_row_idx + 1 if cm.header_end_row_idx >= 0 else 0

    for cells in all_rows[start:]:
        if cm.name_idx >= 0:
            item = _extract_row_mapped(cells, cm)
        else:
            item = _extract_row_heuristic(cells)
        if item:
            rows.append(item)
    return rows


def parse_xlsx(path: Path) -> List[ExtractedRow]:
    rows: List[ExtractedRow] = []
    
    # 1. Try Calamine
    try:
        from python_calamine import CalamineWorkbook
        wb = CalamineWorkbook.from_path(str(path))
        for sheet_name in wb.sheet_names:
            try:
                data = wb.get_sheet_by_name(sheet_name).to_python()
                all_rows = [[str(c).strip() if c is not None else "" for c in r] for r in data if r]
                rows.extend(_extract_from_sheet_rows(all_rows))
            except Exception:
                continue
    except Exception as e:
        logger.debug(f"Calamine parse failed or missing for {path}: {e}")

    if rows:
        return rows

    # 2. Try xlrd for .xls files
    if path.suffix.lower() == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(str(path))
            for sheet in wb.sheets():
                all_rows = []
                for r in range(sheet.nrows):
                    vals = [str(sheet.cell_value(r, c)).strip() for c in range(sheet.ncols)]
                    all_rows.append(vals)
                rows.extend(_extract_from_sheet_rows(all_rows))
        except Exception as e:
            logger.warning(f"xlrd parse failed for {path}: {e}")
        if rows:
            return rows

    # 3. Fallback to openpyxl
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        for sheet in wb.worksheets:
            all_rows = []
            for row in sheet.iter_rows(values_only=True):
                if not row:
                    continue
                all_rows.append([str(c).strip() if c is not None else "" for c in row])
            rows.extend(_extract_from_sheet_rows(all_rows))
    except Exception as e:
        logger.warning(f"openpyxl parse failed for {path}: {e}")

    return rows


# ──────────────── Dispatcher ────────────────

def parse_file(path: Path) -> tuple[str, List[ExtractedRow]]:
    """Return (file_type, rows)."""
    ext = path.suffix.lower().lstrip(".")
    if ext == "pdf":
        rows_pdf = parse_pdf(path)
        if rows_pdf:
            return "pdf", rows_pdf
        rows_ocr = parse_pdf_ocr(path)
        return "scan", rows_ocr
    if ext == "docx":
        return "docx", parse_docx(path)
    if ext in ("xlsx", "xls"):
        return ext, parse_xlsx(path)
    if ext in ("jpg", "jpeg", "png", "webp", "tiff"):
        return ext, parse_image_file(path)
    return ext or "unknown", []


# ──────────────── ZIP archive helpers ────────────────

def fix_zip_filename(filename_str: str, flag_bits: int) -> str:
    filename_str = filename_str.replace("\\", "/")
    if not (flag_bits & 0x800):
        try:
            raw_bytes = filename_str.encode("cp437")
        except UnicodeEncodeError:
            raw_bytes = None

        if raw_bytes:
            try:
                return raw_bytes.decode("utf-8")
            except UnicodeDecodeError:
                pass

            candidates = []
            for enc in ("cp866", "cp1251", "cp437"):
                try:
                    s = raw_bytes.decode(enc)
                    cyrillic = sum(1 for c in s if "\u0400" <= c <= "\u04ff")
                    readable = sum(1 for c in s if c.isalnum() or c in "._- ()[]#")
                    box_drawing = sum(
                        1 for c in s if "\u2500" <= c <= "\u259f" or "\u25a0" <= c <= "\u25ff"
                    )
                    score = cyrillic * 100 + readable - box_drawing * 50
                    candidates.append((score, s))
                except UnicodeDecodeError:
                    pass
            if candidates:
                candidates.sort(key=lambda x: x[0], reverse=True)
                return candidates[0][1]

    if not any("\u0400" <= c <= "\u04ff" for c in filename_str):
        for enc in ("cp1251", "latin1", "cp866"):
            try:
                b = filename_str.encode(enc)
                s = b.decode("utf-8")
                if any("\u0400" <= c <= "\u04ff" for c in s):
                    return s
            except (UnicodeEncodeError, UnicodeDecodeError):
                pass

    return filename_str


def iter_archive(zip_path: Path, extract_to: Path) -> Iterable[Path]:
    import zipfile

    extract_to.mkdir(parents=True, exist_ok=True)
    valid_exts = {"pdf", "docx", "xlsx", "xls", "jpg", "jpeg", "png", "webp", "tiff"}
    extracted_paths: List[Path] = []

    with zipfile.ZipFile(zip_path) as zf:
        for member in zf.infolist():
            if member.is_dir():
                continue

            clean_rel_path = fix_zip_filename(member.filename, member.flag_bits)
            clean_path_obj = Path(clean_rel_path)
            clean_parts = [p for p in clean_path_obj.parts if p not in ("..", ".", "")]
            if not clean_parts:
                continue

            target_file = extract_to.joinpath(*clean_parts)
            target_file.parent.mkdir(parents=True, exist_ok=True)

            with zf.open(member) as source, open(target_file, "wb") as target:
                target.write(source.read())

            if target_file.suffix.lower().lstrip(".") in valid_exts:
                extracted_paths.append(target_file)

    return extracted_paths
