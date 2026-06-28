"""LLM Extraction Service (Fallback mechanism for price list parsing).

Supports Groq API out-of-the-box, structured cleanly so that providers
(Groq, OpenAI, Gemini, local Llama) can be easily swapped or extended in the future.
"""
from __future__ import annotations

import json
import logging
import os
import re
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class LLMExtractedItem:
    code: Optional[str]
    service: str
    price: Optional[float]
    currency: str = "KZT"


def _restore_medical_service_name(name: str) -> str:
    """Auto-restore truncated medical suffixes and clean up formatting."""
    s = name.strip()
    s = s.strip(' "\'\t\r\n_>-')
    if not s:
        return ""
    # Capitalize first letter
    s = s[0].upper() + s[1:]
    # Fix common truncated medical suffixes caused by PDF table column chops
    replacements = [
        (r"эктоми$", "эктомия"),
        (r"стоми$", "стомия"),
        (r"томи$", "томия"),
        (r"графи$", "графия"),
        (r"скопи$", "скопия"),
        (r"метри$", "метрия"),
        (r"терапи$", "терапия"),
        (r"пластик$", "пластика"),
        (r"пункци$", "пункция"),
        (r"биопси$", "биопсия"),
        (r"коагуляци$", "коагуляция"),
        (r"резекци$", "резекция"),
        (r"фиксаци$", "фиксация"),
    ]
    for pattern, repl in replacements:
        s = re.sub(pattern, repl, s, flags=re.IGNORECASE)
    return s


def _is_valid_medical_service_name(name: str) -> bool:
    """Strict validation to filter out junk OCR text, broken word fragments, and footnotes."""
    s = name.strip()
    if len(s) < 6:  # Real medical services are longer than 5 chars
        return False
    if s[0].islower():  # Cutoff starting with lowercase letter (e.g. "ей с контрастом")
        return False
    if s.endswith("-") or s.endswith(",") or (s.endswith(")") and "(" not in s):
        return False

    # Check for junk symbols, OCR noise, and broken lines
    if re.search(r"_{2,}|-{2,}|>{1,}|<|ПЫгГ|CH\.0|C\s*--\s*H", s, flags=re.IGNORECASE):
        return False

    # Check for footnote sentences instead of procedure names
    lower_s = s.lower()
    footnote_prefixes = (
        "у родственниц",
        "в биологическом",
        "биологическом материале",
        "материале методом",
        "оформление справки №",
    )
    if any(lower_s.startswith(p) for p in footnote_prefixes):
        return False

    # Check if string is just random code, numbers, or short gibberish
    if re.match(r"^[A-Z0-9.\s\-_]+$", s):
        return False

    return True


class LLMProvider:
    """Base abstract interface for LLM extraction providers."""

    def extract_json(self, raw_text: str, document_name: str) -> str:
        raise NotImplementedError


class GroqProvider(LLMProvider):
    """Groq API provider implementation using OpenAI-compatible chat completions."""

    def __init__(
        self,
        api_key: Optional[str] = None,
        model: str = "llama-3.3-70b-versatile",
        timeout: float = 45.0,
    ):
        key = api_key or os.getenv("GROQ_API_KEY")
        if not key:
            # Try loading from backend/.env or root .env
            for env_path in (
                Path(__file__).resolve().parents[1] / ".env",
                Path.cwd() / ".env",
                Path.cwd() / "backend" / ".env",
            ):
                if env_path.exists():
                    try:
                        for line in env_path.read_text(encoding="utf-8").splitlines():
                            line = line.strip()
                            if line.startswith("GROQ_API_KEY="):
                                key = line.split("=", 1)[1].strip(' "\'')
                                os.environ["GROQ_API_KEY"] = key
                                break
                    except Exception:  # noqa: BLE001
                        pass
                if key:
                    break

        self.api_key = key
        self.model = model
        self.timeout = timeout
        self.endpoint = "https://api.groq.com/openai/v1/chat/completions"

    def extract_json(self, raw_text: str, document_name: str) -> str:
        if not self.api_key:
            raise ValueError("GROQ_API_KEY is not set in environment variables.")

        system_prompt = (
            "You are an expert medical price list extraction assistant for Kazakhstan medical clinics. "
            "Your objective is to extract complete, grammatically correct medical services along with their REAL numeric prices in KZT.\n\n"
            "CRITICAL RULES FOR SERVICE NAMES AND PRICES:\n"
            "- ONLY extract complete medical procedure/analysis names. Reconstruct words split across lines.\n"
            "- Do NOT extract table column numbers, page numbers, item codes, or footnotes as prices! In KZT price lists, real prices are almost always 300 KZT or higher. Ignore fake prices like 5, 9, 12, 16, 69, 73 KZT.\n"
            "- Capitalize the first letter of each medical service name.\n"
            "- Ignore OCR noise, broken word fragments, table headers, and footnote sentences.\n\n"
            "Return ONLY a valid raw JSON array of objects without markdown code blocks:\n"
            "[\n"
            "  {\n"
            '    "code": "101",\n'
            '    "service": "Электрокардиография",\n'
            '    "price": 6000,\n'
            '    "currency": "KZT"\n'
            "  }\n"
            "]\n\n"
            "Rules:\n"
            "- ONLY extract rows with valid, positive numeric prices (>= 300 KZT).\n"
            "- Extract prices strictly as numbers (integers or floats).\n"
            "- Do not invent services or prices not present in the text."
        )

        user_prompt = f"Document Name: {document_name}\n\nRaw Document Text:\n{raw_text[:12000]}"

        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.1,
            "max_tokens": 4096,
        }

        data_bytes = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            self.endpoint,
            data=data_bytes,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}",
                "User-Agent": "MedArchive-LLM-Fallback/1.0",
            },
            method="POST",
        )

        try:
            with urllib.request.urlopen(req, timeout=self.timeout) as response:
                resp_data = json.loads(response.read().decode("utf-8"))
                return resp_data["choices"][0]["message"]["content"]
        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8", errors="ignore")
            raise RuntimeError(f"Groq API HTTP Error {e.code}: {err_body}") from e
        except Exception as e:
            raise RuntimeError(f"Groq API Request failed: {e}") from e


class LLMExtractionService:
    """Service wrapper for extracting price list items from documents using LLM."""

    def __init__(self, provider: Optional[LLMProvider] = None, max_retries: int = 2):
        self.provider = provider or GroqProvider()
        self.max_retries = max_retries

    def extract_raw_text(self, file_path: Path) -> str:
        """Extract plain text or Markdown-like representation from PDF, DOCX, XLSX, TXT."""
        ext = file_path.suffix.lower().lstrip(".")
        text_chunks: List[str] = []

        try:
            if ext == "pdf":
                import pdfplumber

                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        t = page.extract_text()
                        if t:
                            text_chunks.append(f"--- Page {i+1} ---\n{t}")
            elif ext in ("docx", "doc"):
                import docx

                doc = docx.Document(file_path)
                for p in doc.paragraphs:
                    if p.text.strip():
                        text_chunks.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_txt:
                            text_chunks.append(row_txt)
            elif ext in ("xlsx", "xls"):
                import openpyxl

                wb = openpyxl.load_workbook(file_path, data_only=True)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                        if row_vals:
                            text_chunks.append(" | ".join(row_vals))
            else:
                text_chunks.append(file_path.read_text(errors="ignore"))
        except Exception as e:  # noqa: BLE001
            logger.warning(f"[LLMExtractionService] Error extracting raw text from {file_path.name}: {e}")

        return "\n".join(text_chunks).strip()

    def parse_with_llm(self, file_path: Path) -> tuple[List[LLMExtractedItem], str]:
        """Call LLM fallback with retries, parse JSON, validate, deduplicate, and return structured items and raw text."""
        document_name = file_path.name
        logger.info(f"[LLM Fallback] Triggered for document: {document_name}")
        start_time = time.time()

        raw_text = self.extract_raw_text(file_path)
        if not raw_text or len(raw_text.strip()) < 20:
            logger.warning(f"[LLM Fallback] Document {document_name} contains insufficient extractable text.")
            return [], ""

        response_text = ""
        last_exception: Optional[Exception] = None

        for attempt in range(1, self.max_retries + 2):
            try:
                logger.info(f"[LLM Fallback] Requesting LLM extraction (Attempt {attempt}/{self.max_retries + 1})...")
                response_text = self.provider.extract_json(raw_text, document_name)
                duration = time.time() - start_time
                logger.info(f"[LLM Fallback] Request completed successfully in {duration:.2f}s for {document_name}")
                break
            except Exception as e:  # noqa: BLE001
                last_exception = e
                logger.warning(f"[LLM Fallback] Attempt {attempt} failed: {e}")
                if attempt <= self.max_retries:
                    time.sleep(1.5 * attempt)

        if not response_text:
            logger.error(f"[LLM Fallback] Failed to receive valid LLM response for {document_name}: {last_exception}")
            return [], ""

        items = self._clean_and_parse_json(response_text, document_name)
        return items, response_text

    def _clean_and_parse_json(self, response_text: str, document_name: str) -> List[LLMExtractedItem]:
        cleaned = response_text.strip()
        # Clean markdown ```json code blocks if present
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
        cleaned = cleaned.strip()

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError as e:
            logger.error(f"[LLM Fallback] JSON Decode Error for {document_name}: {e}")
            return []

        if not isinstance(data, list):
            logger.error(f"[LLM Fallback] Expected JSON list from LLM for {document_name}, got {type(data)}.")
            return []

        valid_items: List[LLMExtractedItem] = []
        seen_keys = set()

        for obj in data:
            if not isinstance(obj, dict):
                continue

            service = obj.get("service")
            if not service or not isinstance(service, str) or not service.strip():
                continue

            service_clean = _restore_medical_service_name(service)
            if not _is_valid_medical_service_name(service_clean):
                continue

            code = obj.get("code")
            code_clean = str(code).strip() if code is not None and str(code).strip() != "null" else None

            # Deduplication check by service name + code
            dedup_key = (service_clean.lower(), code_clean)
            if dedup_key in seen_keys:
                continue
            seen_keys.add(dedup_key)

            # Price validation (Strict check: MUST have realistic KZT price >= 300)
            raw_price = obj.get("price")
            price_val: Optional[float] = None
            if raw_price is not None and str(raw_price).strip() != "null":
                try:
                    price_val = float(raw_price)
                    if price_val < 300.0 or price_val > 10_000_000.0:
                        price_val = None
                except (ValueError, TypeError):
                    price_val = None

            # Discard items without valid numeric price or broken text fragments
            if price_val is None or price_val <= 0:
                continue

            currency = str(obj.get("currency") or "KZT").strip()

            valid_items.append(
                LLMExtractedItem(
                    code=code_clean,
                    service=service_clean,
                    price=price_val,
                    currency=currency,
                )
            )

        logger.info(f"[LLM Fallback] Successfully parsed JSON for {document_name}: {len(valid_items)} valid items extracted.")
        return valid_items
